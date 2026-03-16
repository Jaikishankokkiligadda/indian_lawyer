import os
import io
import re
import time
import datetime
import streamlit as st
from pypdf import PdfReader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

try:
    from langchain_huggingface import HuggingFaceEmbeddings
except ImportError:
    from langchain_community.embeddings import HuggingFaceEmbeddings

try:
    from langchain_community.vectorstores import FAISS
except ImportError:
    from langchain.vectorstores import FAISS

from langchain_groq import ChatGroq

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

st.set_page_config(
    page_title="Indian Lawyer – Satyameva Jayate",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Cormorant+Garamond:wght@500;600;700&family=DM+Sans:wght@400;500;600&display=swap');
html, body, [class*="css"] { font-family: 'DM Sans', sans-serif; }
h1, h2, h3 { font-family: 'Cormorant Garamond', serif; }
.stApp { background: #0c0e14; color: #ddd8cc; }
section[data-testid="stSidebar"] { background: #10131c !important; border-right: 1px solid #1e2535; }
.nyaya-header {
    background: linear-gradient(135deg, #13192a 0%, #0c0e14 60%);
    border: 1px solid #bf9b3055; border-radius: 14px;
    padding: 22px 36px; margin-bottom: 20px; position: relative; overflow: hidden;
}
.nyaya-header::before { content: "⚖"; position: absolute; right: -10px; top: -20px; font-size: 140px; opacity: 0.04; pointer-events: none; }
.nyaya-header h1 { color: #bf9b30; font-size: 2rem; margin: 0 0 2px 0; }
.nyaya-header .subtitle { color: #bf9b3099; font-family: 'Cormorant Garamond', serif; font-size: 1rem; font-style: italic; margin: 0 0 4px 0; }
.nyaya-header p { color: #7a8fa6; margin: 0; font-size: 0.85rem; }
.user-bubble { background: #141c2e; border: 1px solid #1e2d46; border-radius: 16px 16px 4px 16px; padding: 14px 18px; margin: 10px 0 10px 15%; color: #c8ddf0; line-height: 1.65; font-size: 0.95rem; }
.assistant-bubble { background: #111a13; border: 1px solid #bf9b3044; border-left: 3px solid #bf9b30; border-radius: 4px 16px 16px 16px; padding: 14px 18px; margin: 10px 15% 10px 0; color: #ddd8cc; line-height: 1.75; font-size: 0.95rem; }
.bubble-role { font-size: 0.72rem; font-weight: 600; letter-spacing: 0.1em; text-transform: uppercase; margin-bottom: 6px; opacity: 0.55; }
.user-role { color: #5a9fd4; } .ai-role { color: #bf9b30; }
.bubble-meta { font-size: 0.68rem; color: #3a4f65; margin-top: 8px; display: flex; gap: 12px; flex-wrap: wrap; }
.source-tag { display: inline-block; background: #0d1f0f; border: 1px solid #2a4d2e; color: #6bbf6e; font-size: 0.7rem; padding: 2px 9px; border-radius: 20px; margin: 4px 3px 0 0; font-family: monospace; }
.doc-card { background: #10131c; border: 1px solid #1e2535; border-radius: 12px; padding: 18px 20px; cursor: pointer; transition: all 0.2s; margin-bottom: 10px; }
.doc-card:hover { border-color: #bf9b3066; background: #13192a; }
.doc-card-title { font-family: 'Cormorant Garamond', serif; font-size: 1.05rem; color: #bf9b30; font-weight: 600; margin-bottom: 4px; }
.doc-card-desc { color: #6a7a8a; font-size: 0.82rem; }
.stTextInput input, .stTextArea textarea, .stSelectbox select { background: #10131c !important; color: #ddd8cc !important; border: 1px solid #1e2535 !important; border-radius: 8px !important; }
.stTextInput input:focus, .stTextArea textarea:focus { border-color: #bf9b30 !important; box-shadow: 0 0 0 1px #bf9b3033 !important; }
.stButton > button { background: linear-gradient(135deg, #bf9b30, #8a6f1e) !important; color: #0c0e14 !important; border: none !important; border-radius: 8px !important; font-weight: 600 !important; }
.stButton > button:hover { opacity: 0.85 !important; }
.ghost-btn > button { background: transparent !important; color: #7a8fa6 !important; border: 1px solid #1e2535 !important; }
.ghost-btn > button:hover { border-color: #bf9b30 !important; color: #bf9b30 !important; }
.status-chip { display: inline-block; padding: 3px 10px; border-radius: 20px; font-size: 0.75rem; font-weight: 600; }
.status-ok { background:#0d1f0f; border:1px solid #2a4d2e; color:#6bbf6e; }
.status-err { background:#1f0d0d; border:1px solid #4d2a2a; color:#bf6b6b; }
.divider { border-color: #1e2535; margin: 16px 0; }
.section-label { font-size: 0.72rem; font-weight: 600; letter-spacing: 0.12em; text-transform: uppercase; color: #3a4f65; margin: 18px 0 8px 0; }
.doc-preview { background: #f5f0e8; color: #1a1a1a; border-radius: 10px; padding: 28px 32px; font-family: 'Times New Roman', serif; font-size: 0.92rem; line-height: 1.8; white-space: pre-wrap; border: 1px solid #d4c9a8; max-height: 500px; overflow-y: auto; }
.empty-state { text-align: center; padding: 56px 0; color: #2a3a4a; }
.empty-state .icon { font-size: 3.5rem; margin-bottom: 12px; }
.empty-state .title { font-family: 'Cormorant Garamond', serif; font-size: 1.3rem; color: #bf9b30; margin-bottom: 8px; }
.empty-state .sub { font-size: 0.88rem; color: #3a4f65; }
.perf-bar-wrap { background: #1e2535; border-radius: 4px; height: 5px; margin-top: 6px; overflow: hidden; }
.perf-bar { height: 100%; border-radius: 4px; background: linear-gradient(90deg, #bf9b30, #6bbf6e); }
.log-row { display: flex; justify-content: space-between; align-items: center; padding: 7px 0; border-bottom: 1px solid #1a2030; font-size: 0.82rem; }
.log-row:last-child { border-bottom: none; }
.log-time { color: #3a4f65; font-family: monospace; font-size: 0.75rem; }
.log-model { color: #bf9b30; }
.log-latency-good { color: #6bbf6e; font-weight: 600; }
.log-latency-ok   { color: #bfa06b; font-weight: 600; }
.log-latency-slow { color: #bf6b6b; font-weight: 600; }
</style>
""", unsafe_allow_html=True)

# ── Constants ──
VECTORSTORE_PATH = "vectorstore"
EMBED_MODEL      = "sentence-transformers/all-MiniLM-L6-v2"
CHUNK_SIZE       = 800
CHUNK_OVERLAP    = 100

GROQ_MODELS = {
    "Llama 3.1 8B (fastest)":   "llama-3.1-8b-instant",
    "Llama 3.3 70B (smartest)": "llama-3.3-70b-versatile",
    "Mixtral 8x7B (balanced)":  "mixtral-8x7b-32768",
    "Gemma 2 9B":               "gemma2-9b-it",
}

DOCUMENT_TEMPLATES = {
    "FIR Draft": {
        "icon":"📋","desc":"First Information Report to file with police",
        "fields":["complainant_name","complainant_address","complainant_phone","incident_date","incident_place","accused_name","incident_description","witnesses","police_station"],
        "labels":{"complainant_name":"Complainant Full Name","complainant_address":"Complainant Address","complainant_phone":"Complainant Phone","incident_date":"Date of Incident","incident_place":"Place of Incident","accused_name":"Accused Name(s)","incident_description":"Describe the Incident in Detail","witnesses":"Witness Names (if any)","police_station":"Police Station Name & District"},
        "textarea_fields":["complainant_address","incident_description"]
    },
    "Legal Notice": {
        "icon":"📜","desc":"Formal legal notice to another party",
        "fields":["sender_name","sender_address","sender_advocate","recipient_name","recipient_address","notice_subject","facts","demand","reply_days"],
        "labels":{"sender_name":"Sender / Client Full Name","sender_address":"Sender Address","sender_advocate":"Advocate Name (optional)","recipient_name":"Recipient / Opposite Party Name","recipient_address":"Recipient Address","notice_subject":"Subject of Notice","facts":"Facts & Circumstances","demand":"Demand / Relief Sought","reply_days":"Days to Reply (e.g. 15, 30)"},
        "textarea_fields":["sender_address","recipient_address","facts","demand"]
    },
    "Bail Application": {
        "icon":"🔓","desc":"Application for regular or anticipatory bail",
        "fields":["applicant_name","applicant_age","applicant_address","fir_number","police_station","sections_charged","arrest_date","court_name","bail_grounds"],
        "labels":{"applicant_name":"Applicant / Accused Full Name","applicant_age":"Age of Applicant","applicant_address":"Applicant's Address","fir_number":"FIR Number","police_station":"Police Station","sections_charged":"IPC/Other Sections Charged","arrest_date":"Date of Arrest","court_name":"Court Name","bail_grounds":"Grounds for Bail"},
        "textarea_fields":["applicant_address","bail_grounds"]
    },
    "Affidavit": {
        "icon":"📝","desc":"General purpose sworn affidavit",
        "fields":["deponent_name","deponent_age","deponent_address","deponent_occupation","affidavit_subject","statement_content","place","affidavit_date"],
        "labels":{"deponent_name":"Deponent Full Name","deponent_age":"Age","deponent_address":"Address","deponent_occupation":"Occupation","affidavit_subject":"Subject / Purpose of Affidavit","statement_content":"Content / Statements to Declare","place":"Place of Execution","affidavit_date":"Date"},
        "textarea_fields":["deponent_address","statement_content"]
    },
    "Rent Agreement": {
        "icon":"🏠","desc":"Residential / commercial rent agreement draft",
        "fields":["landlord_name","landlord_address","tenant_name","tenant_address","property_address","rent_amount","security_deposit","lease_start","lease_duration","special_terms"],
        "labels":{"landlord_name":"Landlord Full Name","landlord_address":"Landlord Address","tenant_name":"Tenant Full Name","tenant_address":"Tenant Address","property_address":"Property / Premises Address","rent_amount":"Monthly Rent (₹)","security_deposit":"Security Deposit (₹)","lease_start":"Lease Start Date","lease_duration":"Lease Duration (e.g. 11 months)","special_terms":"Special Terms & Conditions (optional)"},
        "textarea_fields":["landlord_address","tenant_address","property_address","special_terms"]
    },
    "Consumer Complaint": {
        "icon":"🛒","desc":"Consumer forum complaint against seller/service",
        "fields":["complainant_name","complainant_address","complainant_phone","opposite_party_name","opposite_party_address","purchase_date","product_service","complaint_details","relief_sought","forum_name"],
        "labels":{"complainant_name":"Complainant Full Name","complainant_address":"Complainant Address","complainant_phone":"Phone Number","opposite_party_name":"Opposite Party / Company Name","opposite_party_address":"Opposite Party Address","purchase_date":"Date of Purchase / Service","product_service":"Product / Service Name","complaint_details":"Details of Complaint / Deficiency","relief_sought":"Relief / Compensation Sought","forum_name":"Consumer Forum Name & District"},
        "textarea_fields":["complainant_address","opposite_party_address","complaint_details","relief_sought"]
    },
}

# ── Session state ──
for key, default in {
    "chat_history":[],"vectorstore":None,"active_tab":"chat",
    "selected_doc":None,"doc_form_data":{},"generated_doc_text":"","generated_doc_bytes":None,
    "metrics":{"total_queries":0,"total_doc_generations":0,"total_tokens_est":0,"query_log":[],"llm_times":[],"retrieval_times":[],"model_usage":{}},
}.items():
    if key not in st.session_state:
        st.session_state[key] = default

def avg(lst): return round(sum(lst)/len(lst)) if lst else 0
def estimate_tokens(text): return max(1, len(text)//4)

def log_query(qtype, model, ret_ms, llm_ms, tokens):
    m = st.session_state.metrics
    m["total_queries"] += 1; m["total_tokens_est"] += tokens
    m["llm_times"].append(llm_ms); m["retrieval_times"].append(ret_ms)
    m["model_usage"][model] = m["model_usage"].get(model,0)+1
    m["query_log"].append({"time":datetime.datetime.now().strftime("%H:%M:%S"),"type":qtype,"model":model,"retrieval_ms":ret_ms,"llm_ms":llm_ms,"total_ms":ret_ms+llm_ms,"tokens":tokens})
    if len(m["query_log"])>50: m["query_log"]=m["query_log"][-50:]

def get_api_key():
    try: return st.secrets["GROQ_API_KEY"]
    except: return os.environ.get("GROQ_API_KEY","")

@st.cache_resource(show_spinner=False)
def load_embeddings():
    return HuggingFaceEmbeddings(model_name=EMBED_MODEL, model_kwargs={"device":"cpu"}, encode_kwargs={"batch_size":32,"normalize_embeddings":True})

def load_llm(model_id, api_key):
    return ChatGroq(model=model_id, groq_api_key=api_key, temperature=0.2, max_tokens=1024, streaming=True)

def build_vectorstore(files, embeddings):
    docs=[]
    for uf in files:
        try:
            reader=PdfReader(uf); text="".join(p.extract_text() or "" for p in reader.pages)
            if text.strip(): docs.append(Document(page_content=text,metadata={"source":uf.name}))
        except Exception as e: st.warning(f"⚠️ {uf.name}: {e}")
    if not docs: return None
    chunks=RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE,chunk_overlap=CHUNK_OVERLAP).split_documents(docs)
    vs=FAISS.from_documents(chunks,embeddings); vs.save_local(VECTORSTORE_PATH); return vs

def load_saved_vs(embeddings):
    if os.path.exists(VECTORSTORE_PATH):
        return FAISS.load_local(VECTORSTORE_PATH,embeddings=embeddings,allow_dangerous_deserialization=True)
    return None

def chat_prompt(question, context, history):
    hist_txt=""
    for t in history[-2:]:
        c=t["content"][:400]+"..." if len(t["content"])>400 else t["content"]
        hist_txt+=f"{'User' if t['role']=='user' else 'Assistant'}: {c}\n"
    ctx=context[:1500]+"..." if len(context)>1500 else context
    return f"""You are Indian Lawyer – Satyameva Jayate, an expert Indian legal assistant.

Reply in this format:
### ⚖️ Legal Explanation
[concise plain-language explanation]

### 📋 Relevant Sections & Articles
[key IPC/BNS/Constitutional provisions]

### 🏛️ Relevant Case Laws
[1-2 landmark cases if applicable]

### 💡 Practical Advice
[actionable next steps]

History: {hist_txt}
Context: {ctx}
Question: {question}
Answer:"""

def doc_gen_prompt(doc_type, fields):
    fs="\n".join(f"- {k.replace('_',' ').title()}: {v}" for k,v in fields.items() if v)
    today=datetime.date.today().strftime("%d %B %Y")
    prompts={
        "FIR Draft":f"Draft a formal FIR for Indian police dated {today}. Include header, complainant details, incident narration, applicable IPC/BNS sections, prayer.\nDetails:\n{fs}\n\nFIR:",
        "Legal Notice":f"Draft a formal Legal Notice under Indian law dated {today}. Include LEGAL NOTICE heading, party details, facts, legal grounds, demand with timeline, consequence.\nDetails:\n{fs}\n\nNotice:",
        "Bail Application":f"Draft a Bail Application for Indian court dated {today} u/s 437/439 CrPC. Include court heading, applicant details, FIR details, 5+ bail grounds, prayer.\nDetails:\n{fs}\n\nApplication:",
        "Affidavit":f"Draft a formal Affidavit under Indian law dated {today}. Include heading, deponent details, numbered statements, solemn declaration, verification.\nDetails:\n{fs}\n\nAffidavit:",
        "Rent Agreement":f"Draft a comprehensive Rent Agreement dated {today} under Indian law with 10+ clauses and signature block.\nDetails:\n{fs}\n\nAgreement:",
        "Consumer Complaint":f"Draft a Consumer Complaint under Consumer Protection Act 2019 dated {today}. Include forum heading, parties, jurisdiction, numbered facts, legal grounds, relief, verification.\nDetails:\n{fs}\n\nComplaint:",
    }
    return prompts.get(doc_type, f"Draft a {doc_type} dated {today}:\n{fs}")

def create_docx(title, content):
    doc=DocxDocument(); s=doc.sections[0]
    s.top_margin=Inches(1); s.bottom_margin=Inches(1); s.left_margin=Inches(1.25); s.right_margin=Inches(1.25)
    tp=doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=tp.add_run(title.upper()); r.bold=True; r.font.size=Pt(14); r.font.color.rgb=RGBColor(0x1a,0x1a,0x5e)
    doc.add_paragraph()
    for line in content.split("\n"):
        line=line.strip()
        if not line: doc.add_paragraph(); continue
        para=doc.add_paragraph(); clean=line.lstrip("#").strip()
        if line.startswith("###") or (line.isupper() and len(line)>4):
            r=para.add_run(clean); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=RGBColor(0x1a,0x1a,0x5e)
            para.paragraph_format.space_before=Pt(10)
        elif line.startswith("**") and line.endswith("**"):
            r=para.add_run(line.strip("*")); r.bold=True
        elif line.startswith("-") or line.startswith("•"):
            para.style=doc.styles["List Bullet"]; r=para.add_run(line.lstrip("-•").strip())
        elif re.match(r"^\d+\.",line):
            para.style=doc.styles["List Number"]; r=para.add_run(re.sub(r"^\d+\.\s*","",line))
        else:
            r=para.add_run(clean)
        r.font.size=Pt(11); para.paragraph_format.space_after=Pt(2)
    buf=io.BytesIO(); doc.save(buf); return buf.getvalue()

# ── SIDEBAR ──
with st.sidebar:
    st.markdown("""<div style='padding:12px 0 6px;'>
      <span style='font-family:"Cormorant Garamond",serif;font-size:1.3rem;color:#bf9b30;font-weight:700;'>Indian Lawyer</span><br>
      <span style='font-family:"Cormorant Garamond",serif;font-size:0.9rem;color:#bf9b3099;font-style:italic;'>Satyameva Jayate</span><br>
      <span style='font-size:0.72rem;color:#3a4f65;'>AI Legal Assistant · Powered by Groq</span>
    </div>""", unsafe_allow_html=True)

    st.markdown('<hr class="divider">', unsafe_allow_html=True)
    st.markdown('<div class="section-label">🔑 Groq API Key</div>', unsafe_allow_html=True)

    env_key = get_api_key()
    if env_key:
        st.markdown('<span class="status-chip status-ok">● API key loaded</span>', unsafe_allow_html=True)
        api_key = env_key
    else:
        api_key = st.text_input("Groq API Key", type="password", placeholder="gsk_...", label_visibility="collapsed")
        st.markdown("<div style='font-size:0.78rem;color:#3a4f65'>Get free key at <a href='https://console.groq.com' style='color:#bf9b30'>console.groq.com</a></div>", unsafe_allow_html=True)

    st.markdown('<hr class="divider">', unsafe_allow_html=True)
    st.markdown('<div class="section-label">🤖 Model</div>', unsafe_allow_html=True)
    model_label = st.selectbox("Model", list(GROQ_MODELS.keys()), label_visibility="collapsed")
    model_id    = GROQ_MODELS[model_label]

    st.markdown('<hr class="divider">', unsafe_allow_html=True)
    st.markdown('<div class="section-label">📚 Knowledge Base</div>', unsafe_allow_html=True)
    top_k   = st.slider("Chunks to retrieve", 2, 8, 3)
    kb_mode = st.radio("", ["Upload new PDFs","Load saved index"], label_visibility="collapsed")

    if kb_mode == "Upload new PDFs":
        uploaded = st.file_uploader("", type="pdf", accept_multiple_files=True, label_visibility="collapsed")
        if st.button("⚡ Build Index", use_container_width=True):
            if not uploaded: st.error("Upload at least one PDF.")
            else:
                with st.spinner("Indexing..."):
                    t0=time.time(); emb=load_embeddings(); vs=build_vectorstore(uploaded,emb); ms=round((time.time()-t0)*1000)
                    if vs: st.session_state.vectorstore=vs; st.success(f"✅ Indexed {len(uploaded)} PDF(s) in {ms}ms")
                    else: st.error("No readable text found.")
    else:
        if st.button("📂 Load Saved Index", use_container_width=True):
            with st.spinner("Loading..."):
                emb=load_embeddings(); vs=load_saved_vs(emb)
                if vs: st.session_state.vectorstore=vs; st.success("✅ Loaded!")
                else: st.error("No saved index found.")

    st.markdown('<hr class="divider">', unsafe_allow_html=True)
    kb_ok  = st.session_state.vectorstore is not None
    kb_tag = '<span class="status-chip status-ok">● Ready</span>' if kb_ok else '<span class="status-chip status-err">● Not loaded</span>'
    m      = st.session_state.metrics
    avg_ms = avg(m["llm_times"])
    lc     = "#6bbf6e" if avg_ms<3000 else "#bfa06b" if avg_ms<8000 else "#bf6b6b"
    st.markdown(f"""<div style='font-size:0.8rem;color:#3a4f65;line-height:2.1;'>
      Knowledge Base: {kb_tag}<br>
      Model: <span style='color:#bf9b30'>{model_label}</span><br>
      Queries: <span style='color:#ddd8cc'>{m["total_queries"]}</span><br>
      Avg LLM: <span style='color:{lc}'>{avg_ms/1000:.1f}s</span>
    </div>""", unsafe_allow_html=True)
    st.markdown('<hr class="divider">', unsafe_allow_html=True)
    st.caption("Powered by Groq · Fast · Free tier")

# ── HEADER ──
st.markdown("""<div class="nyaya-header">
  <h1>Indian Lawyer ⚖️</h1>
  <div class="subtitle">Satyameva Jayate — सत्यमेव जयते</div>
  <p>AI-powered Indian Legal Assistant · Chat · Document Generator · RAG · Metrics · Powered by Groq</p>
</div>""", unsafe_allow_html=True)

# ── TABS ──
t1,t2,t3,t4 = st.columns(4)
with t1:
    if st.button("💬 Legal Chatbot",    use_container_width=True): st.session_state.active_tab="chat";    st.rerun()
with t2:
    if st.button("📄 Doc Generator",   use_container_width=True): st.session_state.active_tab="docs";    st.rerun()
with t3:
    if st.button("📌 Quick Reference", use_container_width=True): st.session_state.active_tab="ref";     st.rerun()
with t4:
    if st.button("📊 Metrics",         use_container_width=True): st.session_state.active_tab="metrics"; st.rerun()
st.markdown("---")

if not api_key:
    st.warning("⚠️ Enter your **Groq API key** in the sidebar. Free key at [console.groq.com](https://console.groq.com)")
    st.stop()

# ══ TAB 1: CHAT ══
if st.session_state.active_tab == "chat":
    if not st.session_state.chat_history:
        st.markdown("""<div class="empty-state"><div class="icon">⚖️</div>
          <div class="title">Ask any Indian Legal Question</div>
          <div class="sub">Upload legal PDFs for document-grounded answers,<br>or ask general IPC / Constitutional / procedural questions.</div>
        </div>""", unsafe_allow_html=True)
    else:
        for turn in st.session_state.chat_history:
            if turn["role"]=="user":
                st.markdown(f'<div class="user-bubble"><div class="bubble-role user-role">You</div>{turn["content"]}</div>', unsafe_allow_html=True)
            else:
                srcs="".join(f'<span class="source-tag">📄 {s}</span>' for s in turn.get("sources",[]))
                ch=turn["content"].replace("\n","<br>")
                meta=turn.get("meta",{})
                mh=f'<div class="bubble-meta"><span>⚡ {meta.get("llm_ms",0)/1000:.1f}s</span>{"<span>🔍 "+str(meta.get("retrieval_ms",0))+"ms</span>" if meta.get("retrieval_ms") else ""}<span>~{meta.get("tokens",0)} tokens</span></div>' if meta else ""
                st.markdown(f'<div class="assistant-bubble"><div class="bubble-role ai-role">⚖️ Satyameva Jayate</div>{ch}<div style="margin-top:8px">{srcs}</div>{mh}</div>', unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)
    question = st.text_area("Question", placeholder="e.g. What are my rights if arrested without warrant? What is IPC 420?", height=90, label_visibility="collapsed")
    c1,c2,c3 = st.columns([3,1,1])
    with c1: ask=st.button("⚖️ Ask Satyameva Jayate", use_container_width=True)
    with c2:
        st.markdown('<div class="ghost-btn">', unsafe_allow_html=True)
        clear=st.button("🗑 Clear", use_container_width=True)
        st.markdown('</div>', unsafe_allow_html=True)
    with c3:
        st.markdown('<div class="ghost-btn">', unsafe_allow_html=True)
        export=st.button("📥 Export", use_container_width=True)
        st.markdown('</div>', unsafe_allow_html=True)

    if clear: st.session_state.chat_history=[]; st.rerun()
    if export and st.session_state.chat_history:
        lines=[]
        for t in st.session_state.chat_history:
            r="YOU" if t["role"]=="user" else "SATYAMEVA JAYATE"
            lines.append(f"[{r}]\n{t['content']}\n")
            if t.get("sources"): lines.append(f"Sources: {', '.join(t['sources'])}\n")
            lines.append("-"*60+"\n")
        st.download_button("⬇️ Download Chat","\n".join(lines),"legal_chat.txt","text/plain")

    if ask and question.strip():
        st.session_state.chat_history.append({"role":"user","content":question.strip()})
        sources=[]; retrieval_ms=0; context="No document context. Answer from general knowledge."
        if st.session_state.vectorstore:
            with st.spinner("🔍 Searching legal documents..."):
                t_ret=time.time()
                docs=st.session_state.vectorstore.as_retriever(search_kwargs={"k":top_k}).invoke(question.strip())
                retrieval_ms=round((time.time()-t_ret)*1000)
            sources=list({d.metadata.get("source","unknown") for d in docs})
            context="\n\n".join(f"[{d.metadata.get('source','?')}]\n{d.page_content}" for d in docs)

        prompt=chat_prompt(question.strip(), context, st.session_state.chat_history[:-1])
        llm=load_llm(model_id, api_key)
        resp_area=st.empty(); full_resp=""; t_llm=time.time()
        with st.spinner("⚖️ Thinking..."):
            for chunk in llm.stream(prompt):
                full_resp+=chunk.content; resp_area.markdown(full_resp+" ▌")
        resp_area.markdown(full_resp)
        llm_ms=round((time.time()-t_llm)*1000); tokens=estimate_tokens(full_resp)
        log_query("chat", model_label, retrieval_ms, llm_ms, tokens)
        st.session_state.chat_history.append({"role":"assistant","content":full_resp,"sources":sources,"meta":{"retrieval_ms":retrieval_ms,"llm_ms":llm_ms,"tokens":tokens}})
        st.rerun()

# ══ TAB 2: DOCS ══
elif st.session_state.active_tab == "docs":
    if not st.session_state.selected_doc:
        st.markdown("### 📄 Choose a Document Type")
        cols=st.columns(3)
        for i,(doc_type,meta) in enumerate(DOCUMENT_TEMPLATES.items()):
            with cols[i%3]:
                st.markdown(f'<div class="doc-card"><div class="doc-card-title">{meta["icon"]} {doc_type}</div><div class="doc-card-desc">{meta["desc"]}</div></div>', unsafe_allow_html=True)
                if st.button("Use this template",key=f"sel_{doc_type}",use_container_width=True):
                    st.session_state.selected_doc=doc_type; st.session_state.doc_form_data={}
                    st.session_state.generated_doc_text=""; st.session_state.generated_doc_bytes=None; st.rerun()
    else:
        doc_type=st.session_state.selected_doc; meta=DOCUMENT_TEMPLATES[doc_type]
        cb,ct=st.columns([1,5])
        with cb:
            st.markdown('<div class="ghost-btn">', unsafe_allow_html=True)
            if st.button("← Back"): st.session_state.selected_doc=None; st.rerun()
            st.markdown('</div>', unsafe_allow_html=True)
        with ct: st.markdown(f"### {meta['icon']} {doc_type}")
        st.markdown("---")
        left,right=st.columns([1,1])
        with left:
            st.markdown("**Fill in the details**")
            form_data={}
            for field in meta["fields"]:
                label=meta["labels"].get(field,field.replace("_"," ").title()); key=f"form_{doc_type}_{field}"
                val=st.text_area(label,key=key,height=80) if field in meta.get("textarea_fields",[]) else st.text_input(label,key=key)
                form_data[field]=val
            st.markdown("<br>", unsafe_allow_html=True)
            generate_btn=st.button(f"✨ Generate {doc_type}", use_container_width=True)
        with right:
            st.markdown("**Generated Document**")
            if generate_btn:
                filled={k:v for k,v in form_data.items() if v.strip()}
                if len(filled)<2: st.error("Please fill in at least a few fields.")
                else:
                    with st.spinner(f"✍️ Drafting {doc_type}..."):
                        llm=load_llm(model_id,api_key); t_doc=time.time()
                        resp=llm.invoke(doc_gen_prompt(doc_type,filled)); text=resp.content
                        doc_ms=round((time.time()-t_doc)*1000); tokens=estimate_tokens(text)
                        st.session_state.generated_doc_text=text
                        st.session_state.metrics["total_doc_generations"]+=1
                        log_query("doc_gen",model_label,0,doc_ms,tokens)
                        if DOCX_AVAILABLE:
                            try: st.session_state.generated_doc_bytes=create_docx(doc_type,text)
                            except Exception as e: st.warning(f"DOCX failed: {e}"); st.session_state.generated_doc_bytes=None
            if st.session_state.generated_doc_text:
                st.markdown(f'<div class="doc-preview">{st.session_state.generated_doc_text}</div>', unsafe_allow_html=True)
                st.markdown("<br>", unsafe_allow_html=True)
                dl1,dl2=st.columns(2)
                with dl1: st.download_button("📥 Download .txt",data=st.session_state.generated_doc_text,file_name=f"{doc_type.lower().replace(' ','_')}.txt",mime="text/plain",use_container_width=True)
                with dl2:
                    if st.session_state.generated_doc_bytes: st.download_button("📥 Download .docx",data=st.session_state.generated_doc_bytes,file_name=f"{doc_type.lower().replace(' ','_')}.docx",mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",use_container_width=True)
                    else: st.info("Install `python-docx` for .docx export")
            else:
                st.markdown('<div class="empty-state" style="padding:40px 0"><div class="icon">📝</div><div class="title">Document will appear here</div><div class="sub">Fill in the form and click Generate</div></div>', unsafe_allow_html=True)

# ══ TAB 3: REFERENCE ══
elif st.session_state.active_tab == "ref":
    st.markdown("### 📌 Quick IPC / Legal Reference")
    ref_data={
        "🔴 IPC – Offences Against Person":[("IPC 299 / BNS 100","Culpable Homicide"),("IPC 300 / BNS 101","Murder"),("IPC 302 / BNS 103","Punishment for Murder – Death or Life Imprisonment"),("IPC 304B / BNS 80","Dowry Death – min 7 years"),("IPC 307 / BNS 109","Attempt to Murder"),("IPC 323 / BNS 115","Voluntarily causing hurt – 1 yr / ₹1000 fine"),("IPC 354 / BNS 74","Assault on woman – 1-5 yrs"),("IPC 376 / BNS 64","Rape – 7 yrs to Life")],
        "💰 IPC – Property & Fraud":[("IPC 378 / BNS 303","Theft"),("IPC 383 / BNS 308","Extortion"),("IPC 392 / BNS 309","Robbery"),("IPC 395 / BNS 310","Dacoity – 10 yrs to Life"),("IPC 406 / BNS 316","Criminal Breach of Trust – 3 yrs"),("IPC 420 / BNS 318","Cheating – 7 yrs"),("IPC 468 / BNS 339","Forgery for Cheating – 7 yrs")],
        "🏛️ Constitutional Articles":[("Article 14","Right to Equality before Law"),("Article 19","Freedom of Speech, Assembly, Movement"),("Article 21","Right to Life and Personal Liberty"),("Article 22","Protection against Arbitrary Arrest"),("Article 32","Right to Constitutional Remedies (SC)"),("Article 226","High Court Writ Jurisdiction"),("Article 300A","Right to Property")],
        "⚖️ Key Bail Sections":[("CrPC 436 / BNSS 478","Bail in Bailable Offences – Right, not discretion"),("CrPC 437 / BNSS 480","Bail in Non-Bailable – Magistrate discretion"),("CrPC 438 / BNSS 482","Anticipatory Bail – Sessions Court / HC"),("CrPC 439 / BNSS 483","Special Powers of HC and Sessions Court")],
        "📋 Limitation Periods":[("3 Years","Most civil suits (money, contract, tort)"),("12 Years","Suits on immovable property"),("90 Days","Consumer complaint (extendable)"),("1 Year","Cheque bounce complaint u/s 138 NI Act"),("30 Days","Motor accident claim (can be condoned)")],
    }
    for cat,items in ref_data.items():
        with st.expander(cat,expanded=False):
            for section,desc in items:
                st.markdown(f"<div style='display:flex;gap:12px;padding:7px 0;border-bottom:1px solid #1e2535;'><div style='min-width:180px;color:#bf9b30;font-weight:600;font-size:0.85rem;font-family:monospace'>{section}</div><div style='color:#aaa09a;font-size:0.88rem'>{desc}</div></div>", unsafe_allow_html=True)
    st.markdown("---")
    st.markdown('<div style="color:#3a4f65;font-size:0.8rem;text-align:center;padding:12px">⚠️ AI-generated reference for informational purposes only. Always consult a qualified advocate.</div>', unsafe_allow_html=True)

# ══ TAB 4: METRICS ══
elif st.session_state.active_tab == "metrics":
    m=st.session_state.metrics
    st.markdown("### 📊 Performance Metrics")
    c1,c2,c3,c4,c5,c6=st.columns(6)
    c1.metric("Total Queries",   m["total_queries"])
    c2.metric("Avg LLM",         f"{avg(m['llm_times'])/1000:.1f}s" if m["llm_times"] else "—")
    c3.metric("Avg Retrieval",   f"{avg(m['retrieval_times'])}ms"    if m["retrieval_times"] else "—")
    c4.metric("Avg Total",       f"{avg([l['total_ms'] for l in m['query_log']])/1000:.1f}s" if m["query_log"] else "—")
    c5.metric("~Tokens",         f"{m['total_tokens_est']:,}")
    c6.metric("Docs Generated",  m["total_doc_generations"])
    st.markdown("---")
    if not m["query_log"]:
        st.markdown('<div class="empty-state" style="padding:40px 0"><div class="icon">📊</div><div class="title">No data yet</div><div class="sub">Use the chatbot or doc generator to collect metrics.</div></div>', unsafe_allow_html=True)
    else:
        ca,cb=st.columns(2)
        with ca:
            st.markdown('<div style="background:#10131c;border:1px solid #1e2535;border-radius:12px;padding:20px"><h4 style="font-family:\'Cormorant Garamond\',serif;color:#bf9b30;font-size:1rem;margin:0 0 14px">⏱ Recent Queries</h4>', unsafe_allow_html=True)
            for entry in reversed(m["query_log"][-10:]):
                s=entry["total_ms"]/1000; cls="log-latency-good" if s<3 else "log-latency-ok" if s<8 else "log-latency-slow"
                ret=f" | 🔍 {entry['retrieval_ms']}ms" if entry["retrieval_ms"] else ""
                st.markdown(f'<div class="log-row"><span><span class="log-time">{entry["time"]}</span> &nbsp;<span class="log-model">{entry["model"][:16]}</span> &nbsp;<span style="color:#3a4f65;font-size:0.75rem">[{entry["type"]}]</span></span><span class="{cls}">⚡ {s:.1f}s{ret}</span></div>', unsafe_allow_html=True)
            st.markdown('</div>', unsafe_allow_html=True)
        with cb:
            st.markdown('<div style="background:#10131c;border:1px solid #1e2535;border-radius:12px;padding:20px"><h4 style="font-family:\'Cormorant Garamond\',serif;color:#bf9b30;font-size:1rem;margin:0 0 14px">🤖 Model Usage</h4>', unsafe_allow_html=True)
            total_u=sum(m["model_usage"].values()) or 1
            for mdl,cnt in sorted(m["model_usage"].items(),key=lambda x:-x[1]):
                pct=round(cnt/total_u*100)
                st.markdown(f'<div style="margin-bottom:12px"><div style="display:flex;justify-content:space-between;margin-bottom:4px"><span style="color:#bf9b30;font-size:0.85rem;font-weight:600">{mdl[:30]}</span><span style="color:#ddd8cc;font-size:0.82rem">{cnt} ({pct}%)</span></div><div class="perf-bar-wrap"><div class="perf-bar" style="width:{pct}%"></div></div></div>', unsafe_allow_html=True)
            st.markdown('</div>', unsafe_allow_html=True)
    st.markdown("---")
    st.markdown('<div class="ghost-btn">', unsafe_allow_html=True)
    if st.button("🗑 Reset Metrics"):
        st.session_state.metrics={"total_queries":0,"total_doc_generations":0,"total_tokens_est":0,"query_log":[],"llm_times":[],"retrieval_times":[],"model_usage":{}}; st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)
