import os
import io
import re
import time
import datetime
import streamlit as st
from dotenv import load_dotenv

# ── Load .env ──
load_dotenv()

# ── Optional imports ──
try:
    from pypdf import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    from langchain_core.documents import Document
    LANGCHAIN_AVAILABLE = True
except ImportError:
    LANGCHAIN_AVAILABLE = False

try:
    from langchain_huggingface import HuggingFaceEmbeddings
except ImportError:
    try:
        from langchain_community.embeddings import HuggingFaceEmbeddings
    except ImportError:
        HuggingFaceEmbeddings = None

try:
    from langchain_community.vectorstores import FAISS
except ImportError:
    try:
        from langchain.vectorstores import FAISS
    except ImportError:
        FAISS = None

try:
    from langchain_groq import ChatGroq
    GROQ_AVAILABLE = True
except ImportError:
    GROQ_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ══════════════════════════════════════════════
# PAGE CONFIG
# ══════════════════════════════════════════════
st.set_page_config(
    page_title="Indian Lawyer – Satyameva Jayate",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ══════════════════════════════════════════════
# CUSTOM CSS  (navy + gold aesthetic)
# ══════════════════════════════════════════════
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=EB+Garamond:ital,wght@0,400;0,500;1,400&family=DM+Mono:wght@300;400&display=swap');

html, body, [class*="css"] {
    font-family: 'EB Garamond', Georgia, serif;
    background-color: #080E1C;
    color: #F2EEE6;
}

/* Sidebar */
section[data-testid="stSidebar"] {
    background: #0D1526 !important;
    border-right: 1px solid rgba(201,168,76,0.2);
}

/* Header banner */
.nyaya-header {
    background: linear-gradient(135deg, #0D1526 0%, #080E1C 70%);
    border: 1px solid rgba(201,168,76,0.3);
    border-radius: 10px;
    padding: 20px 32px;
    margin-bottom: 22px;
    position: relative;
    overflow: hidden;
}
.nyaya-header::before {
    content: "⚖";
    position: absolute;
    right: -8px; top: -18px;
    font-size: 130px;
    opacity: 0.04;
    pointer-events: none;
}
.nyaya-header h1  { color: #C9A84C; font-size: 1.85rem; margin: 0 0 3px 0; font-family: 'EB Garamond', serif; }
.nyaya-header .sub { color: rgba(201,168,76,0.7); font-family: 'EB Garamond', serif; font-style:italic; font-size:1rem; margin:0 0 4px; }
.nyaya-header p   { color: #525060; margin: 0; font-size: 0.82rem; font-family: 'DM Mono', monospace; font-weight:300; }

/* Chat bubbles */
.user-bubble {
    background: #111C33;
    border: 1px solid rgba(255,255,255,0.08);
    border-radius: 14px 14px 4px 14px;
    padding: 13px 17px;
    margin: 10px 0 10px 18%;
    color: #A09880;
    line-height: 1.7;
    font-size: 0.96rem;
}
.ai-bubble {
    background: rgba(201,168,76,0.06);
    border: 1px solid rgba(201,168,76,0.22);
    border-left: 2.5px solid #C9A84C;
    border-radius: 4px 14px 14px 4px;
    padding: 13px 17px;
    margin: 10px 18% 10px 0;
    color: #F2EEE6;
    line-height: 1.8;
    font-size: 0.96rem;
}
.bubble-role {
    font-family: 'DM Mono', monospace;
    font-size: 0.6rem;
    font-weight: 300;
    letter-spacing: 0.16em;
    text-transform: uppercase;
    margin-bottom: 6px;
    opacity: 0.55;
}
.user-role  { color: #5A9FD4; }
.ai-role    { color: #C9A84C; }
.bubble-meta { font-family:'DM Mono',monospace; font-size:0.62rem; font-weight:300; color:#525060; margin-top:8px; display:flex; gap:12px; flex-wrap:wrap; }

/* Status chips */
.chip { display:inline-block; padding:3px 10px; border-radius:20px; font-family:'DM Mono',monospace; font-size:0.7rem; font-weight:300; }
.chip-ok  { background:#0d1f0f; border:1px solid #2a4d2e; color:#5AB87A; }
.chip-err { background:#1f0d0d; border:1px solid #4d2a2a; color:#BF6B6B; }
.chip-warn{ background:#1f180d; border:1px solid #4d3d1a; color:#C9A84C; }

/* Doc preview */
.doc-preview {
    background: #111C33;
    border: 1px solid rgba(201,168,76,0.25);
    border-radius: 8px;
    padding: 24px 28px;
    font-family: 'EB Garamond', serif;
    font-size: 0.9rem;
    line-height: 1.85;
    white-space: pre-wrap;
    color: #F2EEE6;
    max-height: 480px;
    overflow-y: auto;
}

/* Reference table */
.ref-section { margin-bottom: 1.8rem; }
.ref-head {
    font-family: 'DM Mono', monospace;
    font-size: 0.6rem;
    font-weight: 300;
    letter-spacing: 0.2em;
    text-transform: uppercase;
    color: #9B7730;
    border-bottom: 1px solid rgba(201,168,76,0.25);
    padding-bottom: 5px;
    margin-bottom: 8px;
}
.ref-row {
    display:flex; gap:12px;
    padding:6px 3px;
    border-bottom:1px solid rgba(255,255,255,0.05);
    font-size:0.88rem;
}
.ref-row:last-child { border-bottom:none; }
.ref-code { min-width:195px; font-family:'DM Mono',monospace; font-size:0.7rem; font-weight:300; color:#C9A84C; }
.ref-desc { color:#A09880; }

/* Metrics */
.metric-card {
    background: #111C33;
    border: 1px solid rgba(201,168,76,0.15);
    border-radius: 8px;
    padding: 14px 18px;
    text-align: center;
}
.metric-val { font-family:'EB Garamond',serif; font-size:1.9rem; color:#C9A84C; line-height:1; margin-bottom:4px; }
.metric-lbl { font-family:'DM Mono',monospace; font-size:0.58rem; font-weight:300; letter-spacing:0.14em; text-transform:uppercase; color:#525060; }
.log-row {
    display:flex; justify-content:space-between; align-items:center;
    padding:6px 0; border-bottom:1px solid rgba(255,255,255,0.05);
    font-family:'DM Mono',monospace; font-size:0.7rem; font-weight:300;
}
.log-row:last-child { border-bottom:none; }
.log-time    { color:#525060; }
.log-model   { color:#C9A84C; }
.lat-good    { color:#5AB87A; font-weight:400; }
.lat-ok      { color:#C9A84C; font-weight:400; }
.lat-slow    { color:#BF6B6B; font-weight:400; }
.perf-track  { background:#162040; border-radius:3px; height:4px; margin-top:5px; overflow:hidden; }
.perf-fill   { height:100%; border-radius:3px; background:linear-gradient(90deg,#9B7730,#5AB87A); }

/* Lawyer cards */
.lawyer-card {
    background:#111C33;
    border:1px solid rgba(255,255,255,0.07);
    border-radius:8px;
    padding:14px 18px;
    margin-bottom:10px;
    transition: border-color 0.2s;
}
.lawyer-card:hover { border-color:rgba(201,168,76,0.3); }
.lc-name  { font-size:0.97rem; font-weight:500; color:#F2EEE6; margin-bottom:1px; }
.lc-role  { font-size:0.8rem; color:#A09880; font-style:italic; margin-bottom:5px; }
.lc-tags  { display:flex; flex-wrap:wrap; gap:5px; }
.lc-tag   { font-family:'DM Mono',monospace; font-size:0.57rem; font-weight:300; color:#9B7730; border:1px solid rgba(201,168,76,0.2); padding:1px 8px; border-radius:20px; }
.lc-stat  { font-family:'EB Garamond',serif; font-size:1rem; color:#C9A84C; }
.lc-sub   { font-family:'DM Mono',monospace; font-size:0.58rem; font-weight:300; color:#525060; }

/* Misc */
.divider-gold { border-color:rgba(201,168,76,0.15); margin:14px 0; }
.section-lbl  { font-family:'DM Mono',monospace; font-size:0.58rem; font-weight:300; letter-spacing:0.18em; text-transform:uppercase; color:#525060; margin:16px 0 7px; }
.stTextInput input, .stTextArea textarea, .stSelectbox select {
    background: #111C33 !important;
    color: #F2EEE6 !important;
    border: 1px solid rgba(255,255,255,0.1) !important;
    border-radius: 6px !important;
    font-family: 'EB Garamond', serif !important;
}
.stTextInput input:focus, .stTextArea textarea:focus {
    border-color: rgba(201,168,76,0.5) !important;
    box-shadow: 0 0 0 1px rgba(201,168,76,0.2) !important;
}
.stButton > button {
    background: linear-gradient(135deg, #C9A84C, #9B7730) !important;
    color: #080E1C !important;
    border: none !important;
    border-radius: 4px !important;
    font-family: 'DM Mono', monospace !important;
    font-size: 0.68rem !important;
    font-weight: 400 !important;
    letter-spacing: 0.1em !important;
    text-transform: uppercase !important;
}
.stButton > button:hover { opacity: 0.88 !important; }
.stDownloadButton > button {
    background: transparent !important;
    color: #A09880 !important;
    border: 1px solid rgba(255,255,255,0.1) !important;
    font-family: 'DM Mono', monospace !important;
    font-size: 0.66rem !important;
}
.stDownloadButton > button:hover { color:#C9A84C !important; border-color:rgba(201,168,76,0.4) !important; }
.stTabs [data-baseweb="tab-list"] { background: #0D1526; border-bottom:1px solid rgba(201,168,76,0.15); }
.stTabs [data-baseweb="tab"] { font-family:'DM Mono',monospace; font-size:0.68rem; font-weight:300; letter-spacing:0.1em; text-transform:uppercase; color:#525060; }
.stTabs [aria-selected="true"] { color:#C9A84C !important; border-bottom:2px solid #C9A84C !important; }
</style>
""", unsafe_allow_html=True)

# ══════════════════════════════════════════════
# CONSTANTS
# ══════════════════════════════════════════════
VECTORSTORE_PATH = "vectorstore"
EMBED_MODEL      = "sentence-transformers/all-MiniLM-L6-v2"
CHUNK_SIZE       = 800
CHUNK_OVERLAP    = 100

GROQ_MODELS = {
    "Llama 3.1 8B  —  Fastest":    "llama-3.1-8b-instant",
    "Llama 3.3 70B  —  Smartest":  "llama-3.3-70b-versatile",
    "Mixtral 8×7B  —  Balanced":   "mixtral-8x7b-32768",
    "Gemma 2 9B":                  "gemma2-9b-it",
}

DOCUMENT_TEMPLATES = {
    "FIR Draft": {
        "icon": "📋", "desc": "First Information Report to file with police",
        "fields": ["complainant_name","complainant_address","complainant_phone",
                   "incident_date","incident_place","accused_name",
                   "incident_description","witnesses","police_station"],
        "labels": {
            "complainant_name":    "Complainant Full Name",
            "complainant_address": "Complainant Address",
            "complainant_phone":   "Complainant Phone",
            "incident_date":       "Date of Incident",
            "incident_place":      "Place of Incident",
            "accused_name":        "Accused Name(s)",
            "incident_description":"Describe the Incident in Detail",
            "witnesses":           "Witness Names (if any)",
            "police_station":      "Police Station Name & District",
        },
        "textarea_fields": ["complainant_address","incident_description"],
        "prompt": lambda fields, today: f"""Draft a formal FIR for Indian police, dated {today}.
Include: station header, FIR number placeholder, complainant details, detailed incident narration, applicable IPC/BNS sections, prayer for action.
Details:
{fields}
Write the complete FIR:""",
    },
    "Legal Notice": {
        "icon": "📜", "desc": "Formal legal notice to another party",
        "fields": ["sender_name","sender_address","sender_advocate","recipient_name",
                   "recipient_address","notice_subject","facts","demand","reply_days"],
        "labels": {
            "sender_name":      "Sender / Client Full Name",
            "sender_address":   "Sender Address",
            "sender_advocate":  "Advocate Name (optional)",
            "recipient_name":   "Recipient / Opposite Party Name",
            "recipient_address":"Recipient Address",
            "notice_subject":   "Subject of Notice",
            "facts":            "Facts & Circumstances",
            "demand":           "Demand / Relief Sought",
            "reply_days":       "Days to Reply (e.g. 15, 30)",
        },
        "textarea_fields": ["sender_address","recipient_address","facts","demand"],
        "prompt": lambda fields, today: f"""Draft a formal Legal Notice under Indian law, dated {today}.
Include: LEGAL NOTICE heading, sender/advocate details, recipient details, numbered facts, legal grounds, specific demand with timeline, consequences of non-compliance.
Details:
{fields}
Write the complete Legal Notice:""",
    },
    "Bail Application": {
        "icon": "🔓", "desc": "Application for regular or anticipatory bail",
        "fields": ["applicant_name","applicant_age","applicant_address","fir_number",
                   "police_station","sections_charged","arrest_date","court_name","bail_grounds"],
        "labels": {
            "applicant_name":    "Applicant / Accused Full Name",
            "applicant_age":     "Age of Applicant",
            "applicant_address": "Applicant's Address",
            "fir_number":        "FIR Number",
            "police_station":    "Police Station",
            "sections_charged":  "IPC/Other Sections Charged",
            "arrest_date":       "Date of Arrest",
            "court_name":        "Court Name",
            "bail_grounds":      "Grounds for Bail",
        },
        "textarea_fields": ["applicant_address","bail_grounds"],
        "prompt": lambda fields, today: f"""Draft a Bail Application under Section 437/439 CrPC (or BNSS equivalent) for an Indian court, dated {today}.
Include: court heading, applicant details, FIR details, at least 5 specific grounds for bail, prayer clause.
Details:
{fields}
Write the complete Bail Application:""",
    },
    "Affidavit": {
        "icon": "📝", "desc": "General purpose sworn affidavit",
        "fields": ["deponent_name","deponent_age","deponent_address","deponent_occupation",
                   "affidavit_subject","statement_content","place","affidavit_date"],
        "labels": {
            "deponent_name":      "Deponent Full Name",
            "deponent_age":       "Age",
            "deponent_address":   "Address",
            "deponent_occupation":"Occupation",
            "affidavit_subject":  "Subject / Purpose of Affidavit",
            "statement_content":  "Content / Statements to Declare",
            "place":              "Place of Execution",
            "affidavit_date":     "Date",
        },
        "textarea_fields": ["deponent_address","statement_content"],
        "prompt": lambda fields, today: f"""Draft a formal Affidavit under Indian law, dated {today}.
Include: heading, court/authority details, deponent details, numbered factual statements, solemn declaration, verification clause, signature block.
Details:
{fields}
Write the complete Affidavit:""",
    },
    "Rent Agreement": {
        "icon": "🏠", "desc": "Residential / commercial rent agreement draft",
        "fields": ["landlord_name","landlord_address","tenant_name","tenant_address",
                   "property_address","rent_amount","security_deposit",
                   "lease_start","lease_duration","special_terms"],
        "labels": {
            "landlord_name":     "Landlord Full Name",
            "landlord_address":  "Landlord Address",
            "tenant_name":       "Tenant Full Name",
            "tenant_address":    "Tenant Address",
            "property_address":  "Property / Premises Address",
            "rent_amount":       "Monthly Rent (₹)",
            "security_deposit":  "Security Deposit (₹)",
            "lease_start":       "Lease Start Date",
            "lease_duration":    "Lease Duration (e.g. 11 months)",
            "special_terms":     "Special Terms & Conditions (optional)",
        },
        "textarea_fields": ["landlord_address","tenant_address","property_address","special_terms"],
        "prompt": lambda fields, today: f"""Draft a comprehensive Rental Agreement under Indian law, dated {today}.
Include: parties, property description, at least 10 detailed clauses covering rent, maintenance, termination, lock-in period, dispute resolution, signature block with witnesses.
Details:
{fields}
Write the complete Rent Agreement:""",
    },
    "Consumer Complaint": {
        "icon": "🛒", "desc": "Consumer forum complaint against seller/service",
        "fields": ["complainant_name","complainant_address","complainant_phone",
                   "opposite_party_name","opposite_party_address","purchase_date",
                   "product_service","complaint_details","relief_sought","forum_name"],
        "labels": {
            "complainant_name":      "Complainant Full Name",
            "complainant_address":   "Complainant Address",
            "complainant_phone":     "Phone Number",
            "opposite_party_name":   "Opposite Party / Company Name",
            "opposite_party_address":"Opposite Party Address",
            "purchase_date":         "Date of Purchase / Service",
            "product_service":       "Product / Service Name",
            "complaint_details":     "Details of Complaint / Deficiency",
            "relief_sought":         "Relief / Compensation Sought",
            "forum_name":            "Consumer Forum Name & District",
        },
        "textarea_fields": ["complainant_address","opposite_party_address","complaint_details","relief_sought"],
        "prompt": lambda fields, today: f"""Draft a Consumer Complaint under the Consumer Protection Act 2019, dated {today}.
Include: forum heading, complainant and opposite party details, jurisdiction, numbered statement of facts, legal grounds, specific relief sought, verification.
Details:
{fields}
Write the complete Consumer Complaint:""",
    },
}

LAWYERS = [
    {"initials":"PS","name":"Adv. Priya Sharma",  "role":"Senior Advocate — Supreme Court of India","city":"Delhi",    "specs":["Criminal Law","Constitutional Law","Bail & Appeals"],    "fee":"₹5,000","rating":"4.9"},
    {"initials":"RM","name":"Adv. Rajesh Menon",  "role":"High Court Advocate — Bombay High Court", "city":"Mumbai",   "specs":["Corporate Law","Contract Disputes","Cheque Bounce"],    "fee":"₹3,500","rating":"4.7"},
    {"initials":"SR","name":"Adv. Sunita Rao",    "role":"Family Court Specialist — Karnataka HC",  "city":"Bangalore","specs":["Divorce","Child Custody","Domestic Violence"],           "fee":"₹2,500","rating":"4.8"},
    {"initials":"VJ","name":"Adv. Vikram Joshi",  "role":"Revenue & Property — Pune District Court","city":"Pune",     "specs":["Property Law","Rent Disputes","Land Acquisition"],       "fee":"₹2,000","rating":"4.6"},
    {"initials":"FK","name":"Adv. Fatima Khan",   "role":"Labour & Employment — Telangana HC",      "city":"Hyderabad","specs":["Labour Law","Employment Disputes","Wrongful Termination"],"fee":"₹2,000","rating":"4.7"},
    {"initials":"DN","name":"Adv. Deepak Nair",   "role":"Consumer & Civil — Madras High Court",    "city":"Chennai",  "specs":["Consumer Protection","Civil Suits","Medical Negligence"],"fee":"₹1,500","rating":"4.5"},
    {"initials":"AG","name":"Adv. Ananya Gupta",  "role":"Tax & Regulatory — Calcutta High Court",  "city":"Kolkata",  "specs":["Income Tax","GST","Customs & Excise"],                   "fee":"₹4,000","rating":"4.8"},
    {"initials":"AS","name":"Adv. Arjun Singh",   "role":"Criminal Defence — Allahabad High Court", "city":"Lucknow",  "specs":["Criminal Defence","NDPS","POCSO"],                       "fee":"₹2,500","rating":"4.6"},
    {"initials":"MP","name":"Adv. Meera Pillai",  "role":"Cyber Law & IPR — Karnataka HC",          "city":"Bangalore","specs":["Cyber Crime","Intellectual Property","Data Privacy"],    "fee":"₹3,000","rating":"4.9"},
    {"initials":"HB","name":"Adv. Harish Bhatia", "role":"RTI & Administrative Law — Rajasthan HC", "city":"Jaipur",   "specs":["RTI","PILs","Administrative Law"],                       "fee":"₹1,500","rating":"4.4"},
]

# ══════════════════════════════════════════════
# SESSION STATE
# ══════════════════════════════════════════════
_defaults = {
    "chat_history":       [],
    "vectorstore":        None,
    "active_tab":         "chat",
    "selected_doc":       None,
    "doc_form_data":      {},
    "generated_doc_text": "",
    "generated_doc_bytes":None,
    "metrics": {
        "total_queries":        0,
        "total_doc_generations":0,
        "total_tokens_est":     0,
        "query_log":            [],
        "llm_times":            [],
        "retrieval_times":      [],
        "model_usage":          {},
    },
}
for k, v in _defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

# ══════════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════════
def avg(lst):
    return round(sum(lst) / len(lst)) if lst else 0

def estimate_tokens(text):
    return max(1, len(text) // 4)

def get_api_key():
    """Reads GROQ_API_KEY from .env / environment only — never from UI."""
    return os.getenv("GROQ_API_KEY", "").strip()

def log_query(qtype, model, ret_ms, llm_ms, tokens):
    m = st.session_state.metrics
    m["total_queries"]    += 1
    m["total_tokens_est"] += tokens
    m["llm_times"].append(llm_ms)
    m["retrieval_times"].append(ret_ms)
    m["model_usage"][model] = m["model_usage"].get(model, 0) + 1
    m["query_log"].append({
        "time":         datetime.datetime.now().strftime("%H:%M:%S"),
        "type":         qtype,
        "model":        model,
        "retrieval_ms": ret_ms,
        "llm_ms":       llm_ms,
        "total_ms":     ret_ms + llm_ms,
        "tokens":       tokens,
    })
    if len(m["query_log"]) > 50:
        m["query_log"] = m["query_log"][-50:]

# ══════════════════════════════════════════════
# CACHED RESOURCES
# ══════════════════════════════════════════════
@st.cache_resource(show_spinner=False)
def load_embeddings():
    if HuggingFaceEmbeddings is None:
        return None
    return HuggingFaceEmbeddings(
        model_name=EMBED_MODEL,
        model_kwargs={"device": "cpu"},
        encode_kwargs={"batch_size": 32, "normalize_embeddings": True},
    )

def load_llm(model_id):
    api_key = get_api_key()
    if not GROQ_AVAILABLE:
        st.error("langchain-groq is not installed. Run: pip install langchain-groq")
        return None
    return ChatGroq(
        model=model_id,
        groq_api_key=api_key,
        temperature=0.2,
        max_tokens=1200,
        streaming=True,
    )

def build_vectorstore(files, embeddings):
    if not PDF_AVAILABLE or not LANGCHAIN_AVAILABLE:
        st.error("Install pypdf and langchain: pip install pypdf langchain-text-splitters")
        return None
    docs = []
    for uf in files:
        try:
            reader = PdfReader(uf)
            text = "".join(p.extract_text() or "" for p in reader.pages)
            if text.strip():
                docs.append(Document(page_content=text, metadata={"source": uf.name}))
        except Exception as e:
            st.warning(f"⚠️ {uf.name}: {e}")
    if not docs:
        return None
    splitter = RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
    chunks   = splitter.split_documents(docs)
    vs = FAISS.from_documents(chunks, embeddings)
    vs.save_local(VECTORSTORE_PATH)
    return vs

def load_saved_vs(embeddings):
    if FAISS is None or not os.path.exists(VECTORSTORE_PATH):
        return None
    return FAISS.load_local(
        VECTORSTORE_PATH,
        embeddings=embeddings,
        allow_dangerous_deserialization=True,
    )

def chat_prompt(question, context, history):
    hist_txt = ""
    for t in history[-4:]:
        c = t["content"][:500] + "…" if len(t["content"]) > 500 else t["content"]
        hist_txt += f"{'User' if t['role'] == 'user' else 'Assistant'}: {c}\n"
    ctx = context[:2000] + "…" if len(context) > 2000 else context
    return f"""You are Indian Lawyer – Satyameva Jayate, a senior Indian legal expert.

Reply using EXACTLY this structure:

### ⚖️ Legal Explanation
[Clear plain-language explanation]

### 📋 Relevant Sections & Articles
[Key IPC/BNS/Constitutional provisions with section numbers]

### 🏛️ Relevant Case Laws
[1–2 landmark Supreme Court or High Court cases if applicable]

### 💡 Practical Next Steps
[Concrete, actionable steps]

History:
{hist_txt}

Context from uploaded documents:
{ctx}

Question: {question}
Answer:"""

def build_doc_prompt(doc_type, fields):
    today  = datetime.date.today().strftime("%d %B %Y")
    meta   = DOCUMENT_TEMPLATES[doc_type]
    fs_str = "\n".join(f"- {k.replace('_',' ').title()}: {v}" for k, v in fields.items() if v)
    return meta["prompt"](fs_str, today)

def create_docx(title, content):
    doc = DocxDocument()
    s = doc.sections[0]
    s.top_margin    = Inches(1)
    s.bottom_margin = Inches(1)
    s.left_margin   = Inches(1.25)
    s.right_margin  = Inches(1.25)

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp.add_run(title.upper())
    r.bold = True; r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x5e)
    doc.add_paragraph()

    for line in content.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph(); continue
        para  = doc.add_paragraph()
        clean = line.lstrip("#").strip()
        if line.startswith("###") or (line.isupper() and len(line) > 4):
            r = para.add_run(clean); r.bold = True
            r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x5e)
            para.paragraph_format.space_before = Pt(10)
        elif line.startswith("**") and line.endswith("**"):
            r = para.add_run(line.strip("*")); r.bold = True; r.font.size = Pt(11)
        elif line.startswith("-") or line.startswith("•"):
            para.style = doc.styles["List Bullet"]
            r = para.add_run(line.lstrip("-•").strip()); r.font.size = Pt(11)
        elif re.match(r"^\d+\.", line):
            para.style = doc.styles["List Number"]
            r = para.add_run(re.sub(r"^\d+\.\s*", "", line)); r.font.size = Pt(11)
        else:
            r = para.add_run(clean); r.font.size = Pt(11)
        para.paragraph_format.space_after = Pt(2)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# ══════════════════════════════════════════════
# HEADER
# ══════════════════════════════════════════════
api_key = get_api_key()
key_status_html = (
    '<span class="chip chip-ok">● API key loaded from .env</span>'
    if api_key else
    '<span class="chip chip-err">● GROQ_API_KEY not found in .env</span>'
)

st.markdown(f"""
<div class="nyaya-header">
  <h1>Indian Lawyer ⚖️</h1>
  <div class="sub">Satyameva Jayate — सत्यमेव जयते</div>
  <p>AI-powered Indian Legal Assistant · Chat · Document Generator · RAG · Metrics · Powered by Groq</p>
  <div style="margin-top:10px">{key_status_html}</div>
</div>
""", unsafe_allow_html=True)

if not api_key:
    st.error(
        "**GROQ_API_KEY not found.** "
        "Create a `.env` file in this directory and add:\n\n"
        "```\nGROQ_API_KEY=gsk_your_key_here\n```\n\n"
        "Get a free key at [console.groq.com](https://console.groq.com)"
    )
    st.stop()

# ══════════════════════════════════════════════
# SIDEBAR
# ══════════════════════════════════════════════
with st.sidebar:
    st.markdown("""
    <div style='padding:10px 0 4px'>
      <div style='font-family:"EB Garamond",serif;font-size:1.25rem;color:#C9A84C;font-weight:500'>Indian Lawyer</div>
      <div style='font-family:"EB Garamond",serif;font-size:0.88rem;color:rgba(201,168,76,0.6);font-style:italic'>Satyameva Jayate</div>
      <div style='font-family:"DM Mono",monospace;font-size:0.62rem;font-weight:300;color:#525060;margin-top:3px'>AI Legal Assistant · Groq</div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown('<hr class="divider-gold">', unsafe_allow_html=True)
    st.markdown('<div class="section-lbl">🤖 Model</div>', unsafe_allow_html=True)
    model_label = st.selectbox("Model", list(GROQ_MODELS.keys()), label_visibility="collapsed")
    model_id    = GROQ_MODELS[model_label]

    st.markdown('<hr class="divider-gold">', unsafe_allow_html=True)
    st.markdown('<div class="section-lbl">📚 Knowledge Base (RAG)</div>', unsafe_allow_html=True)
    top_k   = st.slider("Chunks to retrieve", 2, 8, 3)
    kb_mode = st.radio("", ["Upload PDFs", "Load saved index"], label_visibility="collapsed")

    if kb_mode == "Upload PDFs":
        uploaded = st.file_uploader("", type="pdf", accept_multiple_files=True, label_visibility="collapsed")
        if st.button("⚡ Build Index", use_container_width=True):
            if not uploaded:
                st.error("Upload at least one PDF.")
            else:
                with st.spinner("Indexing PDFs…"):
                    t0  = time.time()
                    emb = load_embeddings()
                    if emb is None:
                        st.error("langchain-huggingface not installed.")
                    else:
                        vs  = build_vectorstore(uploaded, emb)
                        ms  = round((time.time() - t0) * 1000)
                        if vs:
                            st.session_state.vectorstore = vs
                            st.success(f"✅ Indexed {len(uploaded)} PDF(s) in {ms}ms")
                        else:
                            st.error("No readable text found in uploaded PDFs.")
    else:
        if st.button("📂 Load Saved Index", use_container_width=True):
            with st.spinner("Loading…"):
                emb = load_embeddings()
                vs  = load_saved_vs(emb) if emb else None
                if vs:
                    st.session_state.vectorstore = vs
                    st.success("✅ Loaded saved index!")
                else:
                    st.error("No saved index found. Upload PDFs first.")

    st.markdown('<hr class="divider-gold">', unsafe_allow_html=True)
    kb_ok  = st.session_state.vectorstore is not None
    kb_tag = '<span class="chip chip-ok">● Ready</span>' if kb_ok else '<span class="chip chip-warn">● Not loaded</span>'
    m      = st.session_state.metrics
    avg_ms = avg(m["llm_times"])
    lc     = "#5AB87A" if avg_ms < 3000 else "#C9A84C" if avg_ms < 8000 else "#BF6B6B"
    st.markdown(f"""
    <div style='font-family:"DM Mono",monospace;font-size:0.65rem;font-weight:300;line-height:2;color:#525060'>
      KB: {kb_tag}<br>
      Model: <span style='color:#C9A84C'>{model_label.split('—')[0].strip()}</span><br>
      Queries: <span style='color:#F2EEE6'>{m["total_queries"]}</span><br>
      Avg LLM: <span style='color:{lc}'>{avg_ms/1000:.1f}s</span><br>
      ~Tokens: <span style='color:#F2EEE6'>{m["total_tokens_est"]:,}</span>
    </div>
    """, unsafe_allow_html=True)
    st.markdown('<hr class="divider-gold">', unsafe_allow_html=True)
    st.caption("Powered by Groq · Key from .env")

# ══════════════════════════════════════════════
# TABS
# ══════════════════════════════════════════════
tab_chat, tab_docs, tab_lawyers, tab_ref, tab_metrics = st.tabs([
    "💬 Legal Chat",
    "📄 Doc Generator",
    "👨‍⚖️ Find a Lawyer",
    "📌 Reference",
    "📊 Metrics",
])

# ════════════════════════════════════════════════════
# TAB 1 — CHAT
# ════════════════════════════════════════════════════
with tab_chat:
    if not st.session_state.chat_history:
        st.markdown("""
        <div style='text-align:center;padding:48px 0;color:#162040'>
          <div style='font-size:3rem'>⚖️</div>
          <div style='font-family:"EB Garamond",serif;font-size:1.3rem;color:#C9A84C;margin:10px 0 6px'>Ask any Indian legal question</div>
          <div style='font-size:0.88rem;color:#525060'>Upload legal PDFs for document-grounded answers,<br>or ask general IPC / Constitutional / procedural questions.</div>
        </div>
        """, unsafe_allow_html=True)
    else:
        for turn in st.session_state.chat_history:
            if turn["role"] == "user":
                st.markdown(
                    f'<div class="user-bubble"><div class="bubble-role user-role">You</div>{turn["content"]}</div>',
                    unsafe_allow_html=True,
                )
            else:
                ch   = turn["content"].replace("\n", "<br>")
                meta = turn.get("meta", {})
                meta_html = ""
                if meta:
                    ret = f'<span>🔍 {meta["retrieval_ms"]}ms</span>' if meta.get("retrieval_ms") else ""
                    meta_html = f'<div class="bubble-meta"><span>⚡ {meta["llm_ms"]/1000:.1f}s</span>{ret}<span>~{meta["tokens"]} tokens</span></div>'
                srcs = "".join(
                    f'<span style="font-family:\'DM Mono\',monospace;font-size:0.65rem;font-weight:300;'
                    f'color:#5AB87A;border:1px solid rgba(91,184,122,0.25);padding:1px 8px;'
                    f'border-radius:20px;margin-right:5px">📄 {s}</span>'
                    for s in turn.get("sources", [])
                )
                st.markdown(
                    f'<div class="ai-bubble"><div class="bubble-role ai-role">⚖ Satyameva Jayate</div>'
                    f'{ch}<div style="margin-top:8px">{srcs}</div>{meta_html}</div>',
                    unsafe_allow_html=True,
                )

    question = st.text_area(
        "Your question",
        placeholder="e.g. What are my rights if arrested without warrant?  |  What is IPC 420?",
        height=90,
        label_visibility="collapsed",
    )

    c1, c2, c3 = st.columns([3, 1, 1])
    with c1: ask    = st.button("⚖️ Ask Satyameva Jayate", use_container_width=True)
    with c2: clear  = st.button("🗑 Clear", use_container_width=True)
    with c3: export = st.button("📥 Export chat", use_container_width=True)

    if clear:
        st.session_state.chat_history = []
        st.rerun()

    if export and st.session_state.chat_history:
        lines = []
        for t in st.session_state.chat_history:
            r = "YOU" if t["role"] == "user" else "SATYAMEVA JAYATE"
            lines.append(f"[{r}]\n{t['content']}\n{'─'*60}")
            if t.get("sources"):
                lines.append(f"Sources: {', '.join(t['sources'])}")
        st.download_button(
            "⬇️ Download conversation",
            "\n\n".join(lines),
            f"legal_chat_{datetime.date.today()}.txt",
            "text/plain",
        )

    if ask and question.strip():
        st.session_state.chat_history.append({"role": "user", "content": question.strip()})

        sources        = []
        retrieval_ms   = 0
        context        = "No document context loaded. Answer from general Indian legal knowledge."

        if st.session_state.vectorstore:
            with st.spinner("🔍 Searching legal documents…"):
                t_ret  = time.time()
                docs   = st.session_state.vectorstore.as_retriever(
                    search_kwargs={"k": top_k}
                ).invoke(question.strip())
                retrieval_ms = round((time.time() - t_ret) * 1000)
            sources = list({d.metadata.get("source", "unknown") for d in docs})
            context = "\n\n".join(
                f"[{d.metadata.get('source','?')}]\n{d.page_content}" for d in docs
            )

        prompt = chat_prompt(question.strip(), context, st.session_state.chat_history[:-1])
        llm    = load_llm(model_id)
        if llm is None:
            st.stop()

        resp_area  = st.empty()
        full_resp  = ""
        t_llm      = time.time()

        with st.spinner("⚖️ Consulting…"):
            for chunk in llm.stream(prompt):
                full_resp += chunk.content
                resp_area.markdown(full_resp + " ▌")

        resp_area.markdown(full_resp)
        llm_ms = round((time.time() - t_llm) * 1000)
        tokens = estimate_tokens(full_resp)
        log_query("chat", model_label, retrieval_ms, llm_ms, tokens)

        st.session_state.chat_history.append({
            "role":    "assistant",
            "content": full_resp,
            "sources": sources,
            "meta":    {"retrieval_ms": retrieval_ms, "llm_ms": llm_ms, "tokens": tokens},
        })
        st.rerun()

# ════════════════════════════════════════════════════
# TAB 2 — DOC GENERATOR
# ════════════════════════════════════════════════════
with tab_docs:
    if not st.session_state.selected_doc:
        st.markdown("### 📄 Choose a Document Type")
        cols = st.columns(3)
        for i, (doc_type, meta) in enumerate(DOCUMENT_TEMPLATES.items()):
            with cols[i % 3]:
                st.markdown(
                    f'<div style="background:#111C33;border:1px solid rgba(201,168,76,0.15);border-radius:8px;'
                    f'padding:14px 16px;margin-bottom:6px"><div style="font-family:\'EB Garamond\',serif;'
                    f'font-size:1rem;color:#C9A84C;margin-bottom:3px">{meta["icon"]} {doc_type}</div>'
                    f'<div style="font-size:0.78rem;color:#525060;font-style:italic">{meta["desc"]}</div></div>',
                    unsafe_allow_html=True,
                )
                if st.button("Select", key=f"sel_{doc_type}", use_container_width=True):
                    st.session_state.selected_doc       = doc_type
                    st.session_state.doc_form_data      = {}
                    st.session_state.generated_doc_text = ""
                    st.session_state.generated_doc_bytes= None
                    st.rerun()
    else:
        doc_type = st.session_state.selected_doc
        meta     = DOCUMENT_TEMPLATES[doc_type]

        cb, ct = st.columns([1, 6])
        with cb:
            if st.button("← Back"):
                st.session_state.selected_doc = None
                st.rerun()
        with ct:
            st.markdown(f"### {meta['icon']} {doc_type}")

        st.markdown("---")
        left, right = st.columns([1, 1])

        with left:
            st.markdown("**Fill in the details**")
            form_data = {}
            for field in meta["fields"]:
                label = meta["labels"].get(field, field.replace("_", " ").title())
                key   = f"form_{doc_type}_{field}"
                if field in meta.get("textarea_fields", []):
                    val = st.text_area(label, key=key, height=80)
                else:
                    val = st.text_input(label, key=key)
                form_data[field] = val

            generate_btn = st.button(f"✨ Generate {doc_type} with AI", use_container_width=True)

        with right:
            st.markdown("**Generated Document**")
            if generate_btn:
                filled = {k: v for k, v in form_data.items() if v.strip()}
                if len(filled) < 2:
                    st.error("Please fill in at least a few fields first.")
                else:
                    with st.spinner(f"✍️ AI drafting {doc_type}…"):
                        llm    = load_llm(model_id)
                        t_doc  = time.time()
                        resp   = llm.invoke(build_doc_prompt(doc_type, filled))
                        text   = resp.content
                        doc_ms = round((time.time() - t_doc) * 1000)
                        tokens = estimate_tokens(text)
                        st.session_state.generated_doc_text = text
                        st.session_state.metrics["total_doc_generations"] += 1
                        log_query("doc_gen", model_label, 0, doc_ms, tokens)
                        if DOCX_AVAILABLE:
                            try:
                                st.session_state.generated_doc_bytes = create_docx(doc_type, text)
                            except Exception as e:
                                st.warning(f"DOCX generation failed: {e}")
                                st.session_state.generated_doc_bytes = None

            if st.session_state.generated_doc_text:
                st.markdown(
                    f'<div class="doc-preview">{st.session_state.generated_doc_text}</div>',
                    unsafe_allow_html=True,
                )
                st.markdown("<br>", unsafe_allow_html=True)
                dl1, dl2 = st.columns(2)
                with dl1:
                    st.download_button(
                        "📥 Download .txt",
                        data=st.session_state.generated_doc_text,
                        file_name=f"{doc_type.lower().replace(' ','_')}.txt",
                        mime="text/plain",
                        use_container_width=True,
                    )
                with dl2:
                    if st.session_state.generated_doc_bytes:
                        st.download_button(
                            "📥 Download .docx",
                            data=st.session_state.generated_doc_bytes,
                            file_name=f"{doc_type.lower().replace(' ','_')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True,
                        )
                    else:
                        st.info("Install `python-docx` for .docx export")
            else:
                st.markdown("""
                <div style='text-align:center;padding:40px 0;color:#162040'>
                  <div style='font-size:2.5rem'>📝</div>
                  <div style='font-family:"EB Garamond",serif;font-size:1.1rem;color:#C9A84C;margin:8px 0 4px'>Document will appear here</div>
                  <div style='font-size:0.82rem;color:#525060'>Fill the form and click Generate</div>
                </div>
                """, unsafe_allow_html=True)

# ════════════════════════════════════════════════════
# TAB 3 — FIND A LAWYER
# ════════════════════════════════════════════════════
with tab_lawyers:
    st.markdown("### 👨‍⚖️ Advocates Across India")
    filter_q = st.text_input(
        "Search",
        placeholder="Filter by city, specialisation, or keyword…",
        label_visibility="collapsed",
    )
    q = filter_q.lower()
    filtered = [
        l for l in LAWYERS
        if not q or q in l["name"].lower() or q in l["city"].lower()
        or any(q in s.lower() for s in l["specs"]) or q in l["role"].lower()
    ] if q else LAWYERS

    for l in filtered:
        tags_html = "".join(f'<span class="lc-tag">{s}</span>' for s in l["specs"])
        tags_html += f'<span class="lc-tag">{l["city"]}</span>'
        cols = st.columns([0.55, 5, 1.2])
        with cols[0]:
            st.markdown(
                f'<div style="width:40px;height:40px;border-radius:50%;background:#162040;'
                f'border:1px solid rgba(201,168,76,0.25);display:flex;align-items:center;'
                f'justify-content:center;font-family:\'EB Garamond\',serif;font-size:0.75rem;'
                f'color:#C9A84C;margin-top:4px">{l["initials"]}</div>',
                unsafe_allow_html=True,
            )
        with cols[1]:
            st.markdown(
                f'<div class="lc-name">{l["name"]}</div>'
                f'<div class="lc-role">{l["role"]}</div>'
                f'<div class="lc-tags">{tags_html}</div>',
                unsafe_allow_html=True,
            )
        with cols[2]:
            st.markdown(
                f'<div class="lc-stat">{l["rating"]} ★</div>'
                f'<div class="lc-sub">rating</div>'
                f'<div class="lc-stat" style="font-size:.85rem;margin-top:4px">{l["fee"]}</div>'
                f'<div class="lc-sub">consult</div>',
                unsafe_allow_html=True,
            )
        st.markdown('<hr style="border-color:rgba(255,255,255,0.05);margin:6px 0">', unsafe_allow_html=True)

# ════════════════════════════════════════════════════
# TAB 4 — REFERENCE
# ════════════════════════════════════════════════════
with tab_ref:
    st.markdown("### 📌 Quick IPC / Legal Reference")

    ref_data = {
        "🔴 IPC — Offences Against Person": [
            ("IPC 299 / BNS 100","Culpable Homicide"),
            ("IPC 300 / BNS 101","Murder"),
            ("IPC 302 / BNS 103","Punishment — death or life imprisonment"),
            ("IPC 304B / BNS 80","Dowry Death — minimum 7 years"),
            ("IPC 307 / BNS 109","Attempt to Murder"),
            ("IPC 323 / BNS 115","Voluntarily causing hurt — 1 yr / ₹1,000 fine"),
            ("IPC 354 / BNS 74","Assault on woman — 1–5 years"),
            ("IPC 376 / BNS 64","Rape — 7 years to Life"),
        ],
        "💰 IPC — Property & Fraud": [
            ("IPC 378 / BNS 303","Theft"),
            ("IPC 383 / BNS 308","Extortion"),
            ("IPC 392 / BNS 309","Robbery"),
            ("IPC 395 / BNS 310","Dacoity — 10 years to Life"),
            ("IPC 406 / BNS 316","Criminal Breach of Trust — 3 years"),
            ("IPC 420 / BNS 318","Cheating — 7 years"),
            ("IPC 468 / BNS 339","Forgery for Cheating — 7 years"),
        ],
        "🏛️ Constitutional Articles": [
            ("Article 14","Right to Equality before Law"),
            ("Article 19","Freedom of Speech, Assembly, Movement"),
            ("Article 21","Right to Life and Personal Liberty"),
            ("Article 22","Protection against Arbitrary Arrest"),
            ("Article 32","Right to Constitutional Remedies (Supreme Court)"),
            ("Article 226","High Court Writ Jurisdiction"),
            ("Article 300A","Right to Property"),
        ],
        "⚖️ Bail Provisions": [
            ("CrPC 436 / BNSS 478","Bail in Bailable Offences — right, not discretion"),
            ("CrPC 437 / BNSS 480","Non-Bailable — Magistrate's discretion"),
            ("CrPC 438 / BNSS 482","Anticipatory Bail — Sessions Court / High Court"),
            ("CrPC 439 / BNSS 483","Special Powers of High Court and Sessions Court"),
        ],
        "📋 Limitation Periods": [
            ("3 Years","Most civil suits (money, contract, tort)"),
            ("12 Years","Suits on immovable property"),
            ("90 Days","Consumer complaint (extendable)"),
            ("1 Year","Cheque bounce u/s 138 NI Act"),
            ("30 Days","Motor accident claim (condonable)"),
        ],
    }

    for cat, items in ref_data.items():
        with st.expander(cat, expanded=False):
            rows = "".join(
                f'<div class="ref-row"><span class="ref-code">{sec}</span>'
                f'<span class="ref-desc">{desc}</span></div>'
                for sec, desc in items
            )
            st.markdown(f'<div>{rows}</div>', unsafe_allow_html=True)

    st.markdown(
        '<div style="color:#525060;font-family:\'DM Mono\',monospace;font-size:0.72rem;'
        'font-weight:300;text-align:center;padding:14px 0">⚠ AI-generated reference for informational '
        'purposes only. Always consult a qualified advocate.</div>',
        unsafe_allow_html=True,
    )

# ════════════════════════════════════════════════════
# TAB 5 — METRICS
# ════════════════════════════════════════════════════
with tab_metrics:
    m = st.session_state.metrics
    st.markdown("### 📊 Session Performance Metrics")

    c1, c2, c3, c4, c5, c6 = st.columns(6)
    for col, val, lbl in [
        (c1, m["total_queries"],               "Total Queries"),
        (c2, f"{avg(m['llm_times'])/1000:.1f}s" if m["llm_times"] else "—", "Avg LLM"),
        (c3, f"{avg(m['retrieval_times'])}ms"  if m["retrieval_times"] else "—", "Avg Retrieval"),
        (c4, f"{avg([e['total_ms'] for e in m['query_log']])/1000:.1f}s" if m["query_log"] else "—", "Avg Total"),
        (c5, f"{m['total_tokens_est']:,}",     "~Tokens"),
        (c6, m["total_doc_generations"],       "Docs Generated"),
    ]:
        col.markdown(
            f'<div class="metric-card"><div class="metric-val">{val}</div>'
            f'<div class="metric-lbl">{lbl}</div></div>',
            unsafe_allow_html=True,
        )

    st.markdown("---")

    if not m["query_log"]:
        st.markdown("""
        <div style='text-align:center;padding:40px 0;color:#162040'>
          <div style='font-size:2.5rem'>📊</div>
          <div style='font-family:"EB Garamond",serif;font-size:1.1rem;color:#C9A84C;margin:8px 0 4px'>No data yet</div>
          <div style='font-size:0.82rem;color:#525060'>Use the chatbot or document generator to collect metrics.</div>
        </div>
        """, unsafe_allow_html=True)
    else:
        ca, cb = st.columns(2)

        with ca:
            st.markdown(
                '<div style="background:#111C33;border:1px solid rgba(201,168,76,0.12);'
                'border-radius:8px;padding:18px"><div style="font-family:\'DM Mono\',monospace;'
                'font-size:0.6rem;font-weight:300;letter-spacing:0.16em;text-transform:uppercase;'
                'color:#9B7730;margin-bottom:12px">⏱ Recent Queries</div>',
                unsafe_allow_html=True,
            )
            for entry in list(reversed(m["query_log"]))[:12]:
                s   = entry["total_ms"] / 1000
                cls = "lat-good" if s < 3 else "lat-ok" if s < 8 else "lat-slow"
                ret = f' · 🔍 {entry["retrieval_ms"]}ms' if entry["retrieval_ms"] else ""
                short_model = entry["model"][:22]
                st.markdown(
                    f'<div class="log-row">'
                    f'<span><span class="log-time">{entry["time"]}</span> &nbsp;'
                    f'<span class="log-model">{short_model}</span> &nbsp;'
                    f'<span style="color:#525060;font-size:0.65rem">[{entry["type"]}]</span></span>'
                    f'<span class="{cls}">⚡ {s:.1f}s{ret}</span></div>',
                    unsafe_allow_html=True,
                )
            st.markdown('</div>', unsafe_allow_html=True)

        with cb:
            st.markdown(
                '<div style="background:#111C33;border:1px solid rgba(201,168,76,0.12);'
                'border-radius:8px;padding:18px"><div style="font-family:\'DM Mono\',monospace;'
                'font-size:0.6rem;font-weight:300;letter-spacing:0.16em;text-transform:uppercase;'
                'color:#9B7730;margin-bottom:12px">🤖 Model Usage</div>',
                unsafe_allow_html=True,
            )
            total_u = sum(m["model_usage"].values()) or 1
            for mdl, cnt in sorted(m["model_usage"].items(), key=lambda x: -x[1]):
                pct = round(cnt / total_u * 100)
                st.markdown(
                    f'<div style="margin-bottom:11px">'
                    f'<div style="display:flex;justify-content:space-between;margin-bottom:4px">'
                    f'<span style="font-family:\'DM Mono\',monospace;font-size:0.65rem;font-weight:300;'
                    f'color:#C9A84C">{mdl[:32]}</span>'
                    f'<span style="font-family:\'DM Mono\',monospace;font-size:0.63rem;font-weight:300;'
                    f'color:#F2EEE6">{cnt} ({pct}%)</span></div>'
                    f'<div class="perf-track"><div class="perf-fill" style="width:{pct}%"></div></div>'
                    f'</div>',
                    unsafe_allow_html=True,
                )
            st.markdown('</div>', unsafe_allow_html=True)

    st.markdown("---")
    if st.button("🗑 Reset Metrics"):
        st.session_state.metrics = {
            "total_queries": 0, "total_doc_generations": 0, "total_tokens_est": 0,
            "query_log": [], "llm_times": [], "retrieval_times": [], "model_usage": {},
        }
        st.rerun()
